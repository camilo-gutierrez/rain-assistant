"""Document file parsers for Rain RAG system.

Extracts plain text from supported file formats.
Some formats require optional dependencies (pypdf, python-docx, etc.).
"""

import csv
import io
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Supported formats ──────────────────────────────────────────────────
# Core formats (no extra deps)
_CORE_EXTENSIONS = {".txt", ".md", ".csv", ".tsv", ".json"}
# Formats requiring optional deps
_OPTIONAL_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".epub"}
# Source code files (syntax-aware chunking handled downstream)
_CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx",
    ".java", ".go", ".rs", ".c", ".cpp", ".h", ".hpp",
    ".rb", ".php", ".swift", ".kt", ".scala",
    ".sh", ".bash", ".zsh", ".ps1",
    ".sql", ".r", ".m", ".lua",
    ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".xml", ".xsl",
}

SUPPORTED_EXTENSIONS = _CORE_EXTENSIONS | _OPTIONAL_EXTENSIONS | _CODE_EXTENSIONS

MAX_FILE_SIZE_MB = 50
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def parse_file(file_path: str) -> str:
    """Parse a file and return its text content.

    Supported formats:
    - .txt, .md: read as UTF-8
    - .pdf: extract text via pypdf
    - .docx: extract text via python-docx
    - .csv, .tsv: tabular data as markdown table
    - .html, .htm: strip tags via BeautifulSoup
    - .json: pretty-printed JSON
    - .epub: chapter text via ebooklib
    - Source code (.py, .js, .ts, etc.): read as UTF-8 with file header

    Raises:
        FileNotFoundError: if the file does not exist.
        ValueError: if the file format is not supported or file exceeds size limit.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    file_size = path.stat().st_size
    if file_size > MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        raise ValueError(
            f"File too large: {size_mb:.1f} MB. "
            f"Maximum allowed size is {MAX_FILE_SIZE_MB} MB."
        )

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    parser = _PARSERS.get(ext)
    if parser:
        return parser(path)
    return ""


# ── Individual parsers ─────────────────────────────────────────────────

def _parse_text(path: Path) -> str:
    """Read a plain text or markdown file."""
    return path.read_text(encoding="utf-8", errors="replace")


MAX_PDF_PAGES = 500
MAX_EXTRACTED_TEXT_BYTES = 10 * 1024 * 1024  # 10MB of extracted text


def _parse_pdf(path: Path) -> str:
    """Parse PDF file with safety limits."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "pypdf is required for PDF support. "
            "Install with: pip install rain-assistant[memory]"
        )

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {type(e).__name__}")

    pages = []
    total_size = 0
    for i, page in enumerate(reader.pages):
        if i >= MAX_PDF_PAGES:
            logger.warning(
                "PDF '%s' has more than %d pages; truncating.", path.name, MAX_PDF_PAGES
            )
            pages.append(
                f"\n[WARNING: PDF truncated at {MAX_PDF_PAGES} pages out of {len(reader.pages)} total]"
            )
            break
        try:
            text = page.extract_text()
        except Exception:
            continue
        if not text:
            continue
        total_size += len(text.encode("utf-8", errors="replace"))
        if total_size > MAX_EXTRACTED_TEXT_BYTES:
            pages.append("[... truncated: extracted text exceeds 10MB limit]")
            break
        # Add page marker for downstream chunking
        pages.append(f"[Page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    """Extract text from .docx files (paragraphs + tables)."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx support. "
            "Install with: pip install python-docx"
        )

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Preserve heading styles as markdown headers
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", "").replace("Heading", "1"))
            except ValueError:
                level = 1
            level = min(level, 6)
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        table_rows = _extract_docx_table(table)
        if table_rows:
            parts.append(table_rows)

    return "\n\n".join(parts)


def _extract_docx_table(table) -> str:
    """Convert a docx table to markdown table format."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    lines = []
    # Header row
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data rows
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _parse_csv(path: Path) -> str:
    """Parse CSV/TSV into markdown table format."""
    text = path.read_text(encoding="utf-8", errors="replace")
    delimiter = "\t" if path.suffix.lower() in (".tsv",) else ","

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)

    if not rows:
        return ""

    # Limit rows to prevent huge outputs
    max_rows = 1000
    truncated = len(rows) > max_rows
    if truncated:
        rows = rows[:max_rows]

    lines = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data
    for row in rows[1:]:
        # Pad or truncate to match header column count
        padded = row + [""] * (len(rows[0]) - len(row))
        lines.append("| " + " | ".join(padded[:len(rows[0])]) + " |")

    if truncated:
        lines.append(f"\n[... truncated at {max_rows} rows]")

    return "\n".join(lines)


def _parse_html(path: Path) -> str:
    """Strip HTML tags and extract text content."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError(
            "beautifulsoup4 is required for HTML support. "
            "Install with: pip install beautifulsoup4 lxml"
        )

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    # Convert headings to markdown
    for level in range(1, 7):
        for heading in soup.find_all(f"h{level}"):
            heading.replace_with(f"\n{'#' * level} {heading.get_text().strip()}\n")

    # Convert lists to markdown
    for li in soup.find_all("li"):
        li.replace_with(f"\n- {li.get_text().strip()}")

    text = soup.get_text(separator="\n")
    # Clean up excessive whitespace
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(line for line in lines if line)


def _parse_json(path: Path) -> str:
    """Pretty-print JSON for readability."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        data = json.loads(raw)
        return json.dumps(data, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        # Return raw text if invalid JSON
        return raw


def _parse_epub(path: Path) -> str:
    """Extract chapter text from EPUB files."""
    try:
        import ebooklib
        from ebooklib import epub
    except ImportError:
        raise ImportError(
            "ebooklib is required for EPUB support. "
            "Install with: pip install ebooklib lxml"
        )

    book = epub.read_epub(str(path), options={"ignore_ncx": True})
    chapters: list[str] = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Without BS4, just strip tags naively
            content = item.get_content().decode("utf-8", errors="replace")
            chapters.append(content)
            continue

        html = item.get_content().decode("utf-8", errors="replace")
        soup = BeautifulSoup(html, "html.parser")

        for tag in soup(["script", "style"]):
            tag.decompose()

        text = soup.get_text(separator="\n")
        text = text.strip()
        if text:
            chapters.append(text)

    return "\n\n---\n\n".join(chapters)


def _parse_code(path: Path) -> str:
    """Read source code with a file header for context."""
    content = path.read_text(encoding="utf-8", errors="replace")
    header = f"# File: {path.name}\n# Language: {_code_language(path.suffix)}\n\n"
    return header + content


def _code_language(ext: str) -> str:
    """Map file extension to language name."""
    lang_map = {
        ".py": "Python", ".js": "JavaScript", ".ts": "TypeScript",
        ".jsx": "JSX", ".tsx": "TSX", ".java": "Java", ".go": "Go",
        ".rs": "Rust", ".c": "C", ".cpp": "C++", ".h": "C Header",
        ".hpp": "C++ Header", ".rb": "Ruby", ".php": "PHP",
        ".swift": "Swift", ".kt": "Kotlin", ".scala": "Scala",
        ".sh": "Shell", ".bash": "Bash", ".zsh": "Zsh", ".ps1": "PowerShell",
        ".sql": "SQL", ".r": "R", ".m": "MATLAB/Objective-C", ".lua": "Lua",
        ".yaml": "YAML", ".yml": "YAML", ".toml": "TOML",
        ".ini": "INI", ".cfg": "Config", ".xml": "XML", ".xsl": "XSLT",
    }
    return lang_map.get(ext, "Unknown")


# ── Parser dispatch table ──────────────────────────────────────────────
_PARSERS = {
    ".txt": _parse_text,
    ".md": _parse_text,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".csv": _parse_csv,
    ".tsv": _parse_csv,
    ".html": _parse_html,
    ".htm": _parse_html,
    ".json": _parse_json,
    ".epub": _parse_epub,
}
# Register all code extensions
for _ext in _CODE_EXTENSIONS:
    if _ext not in _PARSERS:
        _PARSERS[_ext] = _parse_code
